import server # type: ignore
import re
import json
import os
import time
from google.cloud import vision
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openai
from dotenv import load_dotenv

# Загрузка API-ключей из .env файла
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

def extract_text_with_vision_api(image_path):
    """Распознаёт текст с изображения с помощью Google Vision API."""
    print("\n🔍 Извлекаем текст с изображения...")
    client = vision.ImageAnnotatorClient()
    
    with open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if not texts:
        print("❌ Текст не найден на изображении.")
        return ""

    extracted_text = texts[0].description.strip()
    print(f"✅ Извлечённый текст:\n{extracted_text[:500]}...")
    return extracted_text


def process_with_openai(text):
    """Обрабатывает текст с помощью OpenAI (извлекает паспортные данные)."""
    print("\n📡 Отправляем данные в OpenAI для обработки...")
    prompt = f"""
    You are given the following text extracted from an image of a passport:
    {text}
    
    Please extract the following information in a structured JSON format:
    - Type
    - Code
    - Passport Number
    - Surname
    - Given Names
    - Nationality
    - Date of Birth (format: DD.MM.YYYY)
    - Sex ("Ж" for F, "М" for M)
    - Place of Birth
    - Date of Issue (format: DD.MM.YYYY)
    - Date of Expiry (format: DD.MM.YYYY)
    - Authority (translate to ГОСУДАРСТВЕННЫЙ ДЕПАРТАМЕНТ СОЕДИНЕННЫХ ШТАТОВ)
    - MRZ
    
    If any field is missing, return 'Not found'. 
    Format response as a valid JSON object.
    """
    
    for _ in range(3):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[{"role": "system", "content": "You are a helpful assistant."},
                          {"role": "user", "content": prompt}],
                response_format={"type": "json_object"} 
            )

            structured_data = json.loads(response["choices"][0]["message"]["content"])
            print(f"✅ OpenAI вернул данные:\n{json.dumps(structured_data, indent=4, ensure_ascii=False)}")
            return structured_data
        except Exception as e:
            print(f"⚠️ Ошибка OpenAI: {e}")
            time.sleep(2)
    return {}

def translate_text(text, field_name, target_language="ru"):
    """Переводит текст с помощью OpenAI и удаляет лишние строки."""
    if not text or text == "Not found":
        return text
    
    print(f"\n🌍 Переводим поле {field_name}: {text}")
    prompt = f"""
    Translate the following text into {target_language}:
    {text}
    """
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "system", "content": "You are a professional translator."},
                      {"role": "user", "content": prompt}]
        )
        translated_text = response["choices"][0]["message"]["content"].strip()

        # Удаляем возможные лишние фразы GPT
        # translated_text = re.sub(r'УБЕДИТЕСЬ.*НЕИЗМЕННЫМИ\.', '', translated_text, flags=re.IGNORECASE).strip()
        print(f"✅ Переведено: {translated_text}")
        return translated_text
    except Exception as e:
        print(f"⚠️ Ошибка перевода {field_name}: {e}")
        return text
    
def fill_word_template(template_path, output_path, data):
    """Заполняет шаблон Word данными, выделяя ключевые поля жирным шрифтом и делая все заглавными."""
    print("\n📝 Заполняем шаблон Word...")
    doc = Document(template_path)

    # Ключевые поля, которые должны быть жирными
    bold_fields = ["Фамилия", "Имя", "Имя отца", "№ паспорта", "Дата рождения", "Пол", "Место рождения", "Дата выдачи", "Дата истечения срока", "Орган выдачи"]

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                # Разбиваем текст на части, заменяя плейсхолдер на нужные данные
                new_text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                paragraph.clear()
                
                # Создаём новый run с форматированием
                run = paragraph.add_run(new_text)
                run.font.size = Pt(12)
                # Делаем жирным, если поле в списке bold_fields
                run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run(value.upper())
                            run.font.size = Pt(12)
                            run.bold = True
                                

    doc.save(output_path)
    print(f"✅ Документ сохранён: {output_path}")


def format_mrz(raw_text):
    """
    Форматирует MRZ-зону паспорта в стандартный ICAO-формат.

    - Убирает лишние пробелы
    - Гарантирует две строки по 44 символа
    - Исправляет возможные ошибки OCR
    """
    print("\n📌 Форматируем MRZ...")

    # Удаляем пробелы и оставляем только допустимые символы
    raw_text = re.sub(r"[^A-Z0-9<]", "", raw_text.upper())

    # Если MRZ пришел в одной строке, разбиваем вручную
    if len(raw_text) >= 88:
        mrz_line_1 = raw_text[:44]
        mrz_line_2 = raw_text[44:88]
    else:
        # Разбиваем по << если OpenAI вернул строки корректно
        lines = raw_text.split("\n")
        mrz_lines = [line.strip() for line in lines if "<<" in line or len(line) > 30]

        if len(mrz_lines) < 2:
            print("⚠️ Ошибка: Не удалось найти две строки MRZ! Пробуем разделить автоматически.")
            raw_text = raw_text.ljust(88, "<")  # Добавляем << если длина меньше 88
            mrz_line_1 = raw_text[:44]
            mrz_line_2 = raw_text[44:88]
        else:
            mrz_line_1 = mrz_lines[0].ljust(44, "<")[:44]
            mrz_line_2 = mrz_lines[1].ljust(44, "<")[:44]

    formatted_mrz = f"{mrz_line_1}\n{mrz_line_2}"
    print(f"✅ Отформатированный MRZ:\n{formatted_mrz}")
    return formatted_mrz



def process_document(image_path, template_path, output_path):
    """Основной процесс: OCR -> обработка -> перевод -> заполнение шаблона."""
    try:
        recognized_text = extract_text_with_vision_api(image_path)
        if not recognized_text:
            print("❌ Изображение не содержит текста.")
            return
        
        structured_data = process_with_openai(recognized_text)

        # Форматируем MRZ перед вставкой в документ
        structured_data["MRZ"] = format_mrz(structured_data["MRZ"])

        # 🔹 Переводим ключевые поля перед заполнением Word
        print("\n🌍 Переводим паспортные данные...")
        fields_to_translate = ["Surname", "Given Names", "Nationality", "Place of Birth", "Authority"]
        for field in fields_to_translate:
            if field in structured_data and structured_data[field] != "Not found":
                structured_data[field] = translate_text(structured_data[field], field, "ru")
        
        print("✅ Переведённые данные:")
        print(json.dumps(structured_data, indent=4, ensure_ascii=False))

        # Заполняем шаблон Word
        fill_word_template(template_path, output_path, structured_data)

        print("\n🎉 Процесс завершён успешно!")
    except Exception as e:
        print(f"❌ Ошибка: {e}")


if __name__ == "__main__":
    process_document("input.jpg", "template.docx", "output.docx")
