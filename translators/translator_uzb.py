import re
import json
import os
import time
from google.cloud import vision
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openai
from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

def extract_text_with_vision_api(image_path):
    """–†–∞—Å–ø–æ–∑–Ω–∞—ë—Ç —Ç–µ–∫—Å—Ç —Å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è —Å –ø–æ–º–æ—â—å—é Google Vision API."""
    print("\nüîç –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç —Å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è...")
    client = vision.ImageAnnotatorClient()
    
    with open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if not texts:
        print("‚ùå –¢–µ–∫—Å—Ç –Ω–µ –Ω–∞–π–¥–µ–Ω –Ω–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–∏.")
        return ""

    extracted_text = texts[0].description.strip()
    print(f"‚úÖ –ò–∑–≤–ª–µ—á—ë–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç:\n{extracted_text[:500]}...")
    return extracted_text

def process_with_openai(text):
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Ç–µ–∫—Å—Ç —Å –ø–æ–º–æ—â—å—é OpenAI (–∏–∑–≤–ª–µ–∫–∞–µ—Ç –ø–∞—Å–ø–æ—Ä—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ)."""
    print("\nüì° –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –¥–∞–Ω–Ω—ã–µ –≤ OpenAI –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏...")
    prompt = f"""
    You are processing an Uzbek passport. Return only JSON without any extra explanations.

    - Names (Surname, Given Names) **MUST BE TRANSLATED** to Russian.
    - Geographic names (Place of Birth, Authority) **MUST BE TRANSLATED**.
    - Nationality must be translated if written in Latin script.
    - DO NOT ADD EXTRA COMMENTS OR EXPLANATIONS.

    Extract the following data as JSON:
    - Type
    - Code
    - Passport Number
    - Surname (translate to Russian)
    - Given Names (translate to Russian)
    - Father's Name (translate to Russian)
    - Nationality (translate if in Latin)
    - Date of Birth (dd.mm.yyyy)
    - Sex ("–ñ" for F, "–ú" for M)
    - Authority (if MIA, translate like –ú–ò–ù–ò–°–¢–ï–†–°–¢–í–û –í–ù–£–¢–†–ï–ù–ù–ò–• –î–ï–õ,don't forget about comma after Authority and number of authority)
    - Place of Birth (translate)
    - Date of Issue (dd.mm.yyyy)
    - Date of Expiry (dd.mm.yyyy)
    - MRZ

    Text to process:
    {text}
    """

    for _ in range(3):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[{"role": "system", "content": "You are a professional AI that processes passports."},
                          {"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            structured_data = json.loads(response["choices"][0]["message"]["content"])
            print(f"‚úÖ OpenAI –≤–µ—Ä–Ω—É–ª –¥–∞–Ω–Ω—ã–µ:\n{json.dumps(structured_data, indent=4, ensure_ascii=False)}")
            return structured_data
        except Exception as e:
            print(f"‚ö†Ô∏è –û—à–∏–±–∫–∞ OpenAI: {e}")
            time.sleep(2)
    return {}

def fill_word_template(template_path, output_path, data):
    """–ó–∞–ø–æ–ª–Ω—è–µ—Ç —à–∞–±–ª–æ–Ω Word –¥–∞–Ω–Ω—ã–º–∏, –≤—ã–¥–µ–ª—è—è –∫–ª—é—á–µ–≤—ã–µ –ø–æ–ª—è –∂–∏—Ä–Ω—ã–º —à—Ä–∏—Ñ—Ç–æ–º –∏ –¥–µ–ª–∞—è –≤—Å–µ –∑–∞–≥–ª–∞–≤–Ω—ã–º–∏."""
    print("\nüìù –ó–∞–ø–æ–ª–Ω—è–µ–º —à–∞–±–ª–æ–Ω Word...")
    doc = Document(template_path)

    # –ö–ª—é—á–µ–≤—ã–µ –ø–æ–ª—è, –∫–æ—Ç–æ—Ä—ã–µ –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å –∂–∏—Ä–Ω—ã–º–∏
    bold_fields = ["–§–∞–º–∏–ª–∏—è", "–ò–º—è", "–ò–º—è –æ—Ç—Ü–∞", "‚Ññ –ø–∞—Å–ø–æ—Ä—Ç–∞", "–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è", "–ü–æ–ª", "–ú–µ—Å—Ç–æ —Ä–æ–∂–¥–µ–Ω–∏—è", "–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏", "–î–∞—Ç–∞ –∏—Å—Ç–µ—á–µ–Ω–∏—è —Å—Ä–æ–∫–∞", "–û—Ä–≥–∞–Ω –≤—ã–¥–∞—á–∏"]

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                # –†–∞–∑–±–∏–≤–∞–µ–º —Ç–µ–∫—Å—Ç –Ω–∞ —á–∞—Å—Ç–∏, –∑–∞–º–µ–Ω—è—è –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä –Ω–∞ –Ω—É–∂–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ
                new_text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                paragraph.clear()
                
                # –°–æ–∑–¥–∞—ë–º –Ω–æ–≤—ã–π run —Å —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ–º
                run = paragraph.add_run(new_text)
                run.font.size = Pt(12)
                # –î–µ–ª–∞–µ–º –∂–∏—Ä–Ω—ã–º, –µ—Å–ª–∏ –ø–æ–ª–µ –≤ —Å–ø–∏—Å–∫–µ bold_fields
                run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            # –†–∞–∑–±–∏–≤–∞–µ–º —Ç–µ–∫—Å—Ç –Ω–∞ —á–∞—Å—Ç–∏, –∑–∞–º–µ–Ω—è—è –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä –Ω–∞ –Ω—É–∂–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ
                            new_text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                            paragraph.clear()
                            
                            # –°–æ–∑–¥–∞—ë–º –Ω–æ–≤—ã–π run —Å —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ–º
                            run = paragraph.add_run(new_text)
                            run.font.size = Pt(12)
                            # –î–µ–ª–∞–µ–º –∂–∏—Ä–Ω—ã–º, –µ—Å–ª–∏ –ø–æ–ª–µ –≤ —Å–ø–∏—Å–∫–µ bold_fields
                            run.bold = True
    doc.save(output_path)
    print(f"‚úÖ –î–æ–∫—É–º–µ–Ω—Ç —Å–æ—Ö—Ä–∞–Ω—ë–Ω: {output_path}")



def format_mrz(raw_text):
    """
    –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ—Ç MRZ-–∑–æ–Ω—É –ø–∞—Å–ø–æ—Ä—Ç–∞ –≤ —Å—Ç–∞–Ω–¥–∞—Ä—Ç–Ω—ã–π ICAO-—Ñ–æ—Ä–º–∞—Ç.

    - –£–±–∏—Ä–∞–µ—Ç –ª–∏—à–Ω–∏–µ –ø—Ä–æ–±–µ–ª—ã
    - –ì–∞—Ä–∞–Ω—Ç–∏—Ä—É–µ—Ç –¥–≤–µ —Å—Ç—Ä–æ–∫–∏ –ø–æ 44 —Å–∏–º–≤–æ–ª–∞
    - –ò—Å–ø—Ä–∞–≤–ª—è–µ—Ç –≤–æ–∑–º–æ–∂–Ω—ã–µ –æ—à–∏–±–∫–∏ OCR
    """
    print("\nüìå –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º MRZ...")

    # –£–¥–∞–ª—è–µ–º –ø—Ä–æ–±–µ–ª—ã –∏ –æ—Å—Ç–∞–≤–ª—è–µ–º —Ç–æ–ª—å–∫–æ –¥–æ–ø—É—Å—Ç–∏–º—ã–µ —Å–∏–º–≤–æ–ª—ã
    raw_text = re.sub(r"[^A-Z0-9<]", "", raw_text.upper())

    # –ï—Å–ª–∏ MRZ –ø—Ä–∏—à–µ–ª –≤ –æ–¥–Ω–æ–π —Å—Ç—Ä–æ–∫–µ, —Ä–∞–∑–±–∏–≤–∞–µ–º –≤—Ä—É—á–Ω—É—é
    if len(raw_text) >= 88:
        mrz_line_1 = raw_text[:44]
        mrz_line_2 = raw_text[44:88]
    else:
        # –†–∞–∑–±–∏–≤–∞–µ–º –ø–æ << –µ—Å–ª–∏ OpenAI –≤–µ—Ä–Ω—É–ª —Å—Ç—Ä–æ–∫–∏ –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ
        lines = raw_text.split("\n")
        mrz_lines = [line.strip() for line in lines if "<<" in line or len(line) > 30]

        if len(mrz_lines) < 2:
            print("‚ö†Ô∏è –û—à–∏–±–∫–∞: –ù–µ —É–¥–∞–ª–æ—Å—å –Ω–∞–π—Ç–∏ –¥–≤–µ —Å—Ç—Ä–æ–∫–∏ MRZ! –ü—Ä–æ–±—É–µ–º —Ä–∞–∑–¥–µ–ª–∏—Ç—å –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏.")
            raw_text = raw_text.ljust(88, "<")  # –î–æ–±–∞–≤–ª—è–µ–º << –µ—Å–ª–∏ –¥–ª–∏–Ω–∞ –º–µ–Ω—å—à–µ 88
            mrz_line_1 = raw_text[:44]
            mrz_line_2 = raw_text[44:88]
        else:
            mrz_line_1 = mrz_lines[0].ljust(44, "<")[:44]
            mrz_line_2 = mrz_lines[1].ljust(44, "<")[:44]

    formatted_mrz = f"{mrz_line_1}\n{mrz_line_2}"
    print(f"‚úÖ –û—Ç—Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–π MRZ:\n{formatted_mrz}")
    return formatted_mrz


def process_document(image_path, template_path, output_path):
    """–û—Å–Ω–æ–≤–Ω–æ–π –ø—Ä–æ—Ü–µ—Å—Å: OCR -> –æ–±—Ä–∞–±–æ—Ç–∫–∞ -> –ø–µ—Ä–µ–≤–æ–¥ -> –∑–∞–ø–æ–ª–Ω–µ–Ω–∏–µ —à–∞–±–ª–æ–Ω–∞."""
    try:
        recognized_text = extract_text_with_vision_api(image_path)
        if not recognized_text:
            print("‚ùå –ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç —Ç–µ–∫—Å—Ç–∞.")
            return
        
        structured_data = process_with_openai(recognized_text)

        # –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º MRZ –ø–µ—Ä–µ–¥ –≤—Å—Ç–∞–≤–∫–æ–π –≤ –¥–æ–∫—É–º–µ–Ω—Ç
        structured_data["MRZ"] = format_mrz(structured_data["MRZ"])

        fill_word_template(template_path, output_path, structured_data)

        print("\nüéâ –ü—Ä–æ—Ü–µ—Å—Å –∑–∞–≤–µ—Ä—à—ë–Ω —É—Å–ø–µ—à–Ω–æ!")
    except Exception as e:
        print(f"‚ùå –û—à–∏–±–∫–∞: {e}")

if __name__ == "__main__":
    process_document("input.jpg", "template.docx", "output.docx")
