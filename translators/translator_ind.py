import re
import os
import json
import time
from google.cloud import vision
from docx import Document
from docx.shared import Pt
import openai
from dotenv import load_dotenv

# Загрузка API-ключей
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

def extract_text_with_vision_api(image_path):
    """Распознаёт текст с изображения с помощью Google Vision API."""
    print("\n🔍 Извлекаем текст с изображения...")
    client = vision.ImageAnnotatorClient()
    
    with open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if not texts:
        print("❌ Текст не найден на изображении.")
        return ""

    extracted_text = texts[0].description.strip()
    print(f"✅ Извлечённый текст:\n{extracted_text[:500]}...")
    return extracted_text

def process_with_openai(text):
    """Обрабатывает текст с помощью OpenAI (извлекает паспортные данные)."""
    print("\n📡 Обрабатываем паспорт Индии с OpenAI...")

    prompt = f"""
    You are processing an Indian passport. Return only JSON without any extra explanations.

    Extract the following data as JSON:
    - Type
    - Code
    - Passport Number
    - Surname
    - Given Names
    - Nationality
    - Date of Birth (dd.mm.yyyy)
    - Sex ("Ж" for F, "М" for M)
    - Place of Birth
    - Place of Issue
    - Date of Issue (dd.mm.yyyy)
    - Date of Expiry (dd.mm.yyyy)
    - MRZ

    Text to process:
    {text}
    """

    for _ in range(3):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[{"role": "system", "content": "You are a professional AI that processes passports."},
                          {"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            structured_data = json.loads(response["choices"][0]["message"]["content"])
            structured_data = {key: value.upper() for key, value in structured_data.items()}
            print(f"✅ OpenAI вернул данные:\n{json.dumps(structured_data, indent=4, ensure_ascii=False)}")
            return structured_data
        except Exception as e:
            print(f"⚠️ Ошибка OpenAI: {e}")
            time.sleep(2)
    return {}

def translate_text(text, field_name, target_lang="ru"):
    """Переводит текст с помощью OpenAI."""
    try:
        prompt = f"""
        Translate the following into {target_lang}, keeping the format unchanged:
        {text}
        """

        response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[{"role": "system", "content": "You are an expert AI translator."},
                      {"role": "user", "content": prompt}]
        )
        translated_text = response["choices"][0]["message"]["content"].strip()
        print(f"🔄 Переведено [{field_name}]: {text} → {translated_text}")
        return translated_text
    except Exception as e:
        print(f"⚠️ Ошибка перевода для {field_name}: {e}")
        return text

def format_mrz(raw_text):
    """Форматирует MRZ-зону паспорта в стандартный ICAO-формат."""
    print("\n📌 Форматируем MRZ...")

    raw_text = re.sub(r"[^A-Z0-9<]", "", raw_text.upper())

    if len(raw_text) >= 88:
        mrz_line_1 = raw_text[:44]
        mrz_line_2 = raw_text[44:88]
    else:
        lines = raw_text.split("\n")
        mrz_lines = [line.strip() for line in lines if "<<" in line or len(line) > 30]

        if len(mrz_lines) < 2:
            print("⚠️ Ошибка: Не удалось найти две строки MRZ! Пробуем разделить автоматически.")
            raw_text = raw_text.ljust(88, "<")
            mrz_line_1 = raw_text[:44]
            mrz_line_2 = raw_text[44:88]
        else:
            mrz_line_1 = mrz_lines[0].ljust(44, "<")[:44]
            mrz_line_2 = mrz_lines[1].ljust(44, "<")[:44]

    formatted_mrz = f"{mrz_line_1}\n{mrz_line_2}"
    print(f"✅ Отформатированный MRZ:\n{formatted_mrz}")
    return formatted_mrz

def process_document(image_path, template_path, output_path):
    """Обрабатывает паспорт Индии."""
    try:
        recognized_text = extract_text_with_vision_api(image_path)
        if not recognized_text:
            print("❌ Изображение не содержит текста.")
            return
        
        structured_data = process_with_openai(recognized_text)
        structured_data["MRZ"] = format_mrz(structured_data["MRZ"])

        # 🔹 Переводим ключевые поля перед заполнением Word
        print("\n🌍 Переводим паспортные данные...")
        fields_to_translate = ["Surname", "Given Names", "Nationality", "Place of Birth", "Place of Issue"]
        for field in fields_to_translate:
            if field in structured_data and structured_data[field] != "Not found":
                structured_data[field] = translate_text(structured_data[field], field, "ru")
        
        print("✅ Переведённые данные:")
        print(json.dumps(structured_data, indent=4, ensure_ascii=False))

        # Заполняем шаблон Word
        fill_word_template(template_path, output_path, structured_data)

        print("\n🎉 Процесс завершён успешно!")
    except Exception as e:
        print(f"❌ Ошибка: {e}")

def fill_word_template(template_path, output_path, data):
    """Заполняет шаблон Word данными, выделяя ключевые поля жирным шрифтом и делая все заглавными."""
    print("\n📝 Заполняем шаблон Word...")
    doc = Document(template_path)

    # Ключевые поля, которые должны быть жирными
    bold_fields = ["Фамилия", "Имя", "Имя отца", "№ паспорта", "Дата рождения", "Пол", "Место рождения", "Дата выдачи", "Дата истечения срока", "Орган выдачи"]

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                # Разбиваем текст на части, заменяя плейсхолдер на нужные данные
                new_text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                paragraph.clear()
                
                # Создаём новый run с форматированием
                run = paragraph.add_run(new_text)
                run.font.size = Pt(12)
                # Делаем жирным, если поле в списке bold_fields
                run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            # Разбиваем текст на части, заменяя плейсхолдер на нужные данные
                            new_text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                            paragraph.clear()
                            
                            # Создаём новый run с форматированием
                            run = paragraph.add_run(new_text)
                            run.font.size = Pt(12)
                            # Делаем жирным, если поле в списке bold_fields
                            run.bold = True
    doc.save(output_path)
    print(f"✅ Документ сохранён: {output_path}")