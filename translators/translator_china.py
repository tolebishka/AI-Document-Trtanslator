import re
import json
import os
import time
from google.cloud import vision
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openai
from dotenv import load_dotenv

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

def extract_text_with_vision_api(image_path):
    """Распознаёт текст с изображения с помощью Google Vision API."""
    print("\n🔍 Извлекаем текст с изображения...")
    client = vision.ImageAnnotatorClient()
    
    with open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if not texts:
        print("❌ Текст не найден на изображении.")
        return ""

    extracted_text = texts[0].description.strip()
    print(f"✅ Извлечённый текст:\n{extracted_text[:500]}...")
    return extracted_text

def process_with_openai(text):
    """Обрабатывает текст с помощью OpenAI (извлекает паспортные данные)."""
    print("\n📡 Отправляем данные в OpenAI для обработки...")
    prompt = f"""
    You are processing an Uzbek passport. Return only JSON without any extra explanations.

    - Names (Surname, Given Names) **MUST BE TRANSLATED** to Russian.
    - Geographic names (Place of Birth, Authority) **MUST BE TRANSLATED**.
    - Nationality must be translated if written in Latin script.
    - DO NOT ADD EXTRA COMMENTS OR EXPLANATIONS.

    Extract the following data as JSON:
    - Type
    - Code
    - Passport Number
    - Surname (translate to Russian)
    - Given Names (translate to Russian)
    - Nationality (translate if in Latin)
    - Place of Birth (translate to Russian)
    - Place of Issue (translate to Russian)
    - Date of Birth (dd.mm.yyyy)
    - Sex ("Ж" for F, "М" for M)
    - Date of Issue (dd.mm.yyyy)
    - Date of Expiry (dd.mm.yyyy)
    - MRZ

    Text to process:
    {text}
    """

    for _ in range(3):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[{"role": "system", "content": "You are a professional AI that processes passports."},
                          {"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            structured_data = json.loads(response["choices"][0]["message"]["content"])
            print(f"✅ OpenAI вернул данные:\n{json.dumps(structured_data, indent=4, ensure_ascii=False)}")
            return structured_data
        except Exception as e:
            print(f"⚠️ Ошибка OpenAI: {e}")
            time.sleep(2)
    return {}

def fill_word_template(template_path, output_path, data):
    """Заполняет шаблон Word данными, задавая шрифт Times New Roman."""
    print("\n📝 Заполняем шаблон Word...")
    doc = Document(template_path)

    font_name = "Times New Roman"  # Задайте нужный шрифт

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                for run in paragraph.runs:
                    run.font.name = font_name  # Устанавливаем шрифт
                    run.font.size = Pt(12)  # Размер 12 pt
                    r = run._element
                    run.bold = True
                    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # Для китайских/русских символов

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
                            for run in paragraph.runs:
                                run.font.name = font_name
                                run.font.size = Pt(12)
                                r = run._element
                                r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    doc.save(output_path)
    print(f"✅ Документ сохранён: {output_path}")



def format_mrz(raw_text):
    """
    Форматирует MRZ-зону паспорта в стандартный ICAO-формат.

    - Убирает лишние пробелы
    - Гарантирует две строки по 44 символа
    - Исправляет возможные ошибки OCR
    """
    print("\n📌 Форматируем MRZ...")

    # Удаляем пробелы и оставляем только допустимые символы
    raw_text = re.sub(r"[^A-Z0-9<]", "", raw_text.upper())

    # Если MRZ пришел в одной строке, разбиваем вручную
    if len(raw_text) >= 88:
        mrz_line_1 = raw_text[:44]
        mrz_line_2 = raw_text[44:88]
    else:
        # Разбиваем по << если OpenAI вернул строки корректно
        lines = raw_text.split("\n")
        mrz_lines = [line.strip() for line in lines if "<<" in line or len(line) > 30]

        if len(mrz_lines) < 2:
            print("⚠️ Ошибка: Не удалось найти две строки MRZ! Пробуем разделить автоматически.")
            raw_text = raw_text.ljust(88, "<")  # Добавляем << если длина меньше 88
            mrz_line_1 = raw_text[:44]
            mrz_line_2 = raw_text[44:88]
        else:
            mrz_line_1 = mrz_lines[0].ljust(44, "<")[:44]
            mrz_line_2 = mrz_lines[1].ljust(44, "<")[:44]

    formatted_mrz = f"{mrz_line_1}\n{mrz_line_2}"
    print(f"✅ Отформатированный MRZ:\n{formatted_mrz}")
    return formatted_mrz


def process_document(image_path, template_path, output_path):
    """Основной процесс: OCR -> обработка -> перевод -> заполнение шаблона."""
    try:
        recognized_text = extract_text_with_vision_api(image_path)
        if not recognized_text:
            print("❌ Изображение не содержит текста.")
            return
        
        structured_data = process_with_openai(recognized_text)

        # Форматируем MRZ перед вставкой в документ
        structured_data["MRZ"] = format_mrz(structured_data["MRZ"])

        fill_word_template(template_path, output_path, structured_data)

        print("\n🎉 Процесс завершён успешно!")
    except Exception as e:
        print(f"❌ Ошибка: {e}")

if __name__ == "__main__":
    process_document("input.jpg", "template.docx", "output.docx")
