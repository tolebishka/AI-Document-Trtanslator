# import re
# import json
# import os
# import time
# from google.cloud import vision
# from docx import Document
# import openai
# from dotenv import load_dotenv

# load_dotenv()
# openai.api_key = os.getenv("OPENAI_API_KEY")
# os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

# def extract_text_with_vision_api(image_path):
#     """Распознаёт текст с изображения с помощью Google Vision API."""
#     client = vision.ImageAnnotatorClient()

#     with open(image_path, 'rb') as image_file:
#         content = image_file.read()
#     image = vision.Image(content=content)

#     response = client.text_detection(image=image)
#     texts = response.text_annotations

#     if not texts:
#         print("Текст не найден на изображении.")
#         return ""

#     return texts[0].description

# def process_with_openai(text):
#     """Обработка текста с помощью OpenAI."""
#     for attempt in range(3):
#         try:
#             prompt = f"""
#             You are a professional translator specializing in passport documents. Follow these rules STRICTLY:
#             The passport has been provided in Kazakh language. You need to translate from KZ to RUS
#             If dates already provided in Russian language, return this data
#             1. Names (Surname, Given Names, Father's Name) must remain in their original language. Do NOT translate them.
#             2. Geographic names (Place of Birth, Authority) should be translated into Russian only if they are in Latin script. For example:
#             - "Tashkent" -> "Ташкент"
#             - "Almaty" -> "Алматы"
#             - "Ташкент" -> "Ташкент" (no translation needed)
#             3. Nationality should be translated into Russian only if it is in Latin script. For example:
#             - "Uzbekistan" -> "Узбекистан"
#             - "Қазақстан" -> "Қазақстан" (no translation needed)
#             4. Return ONLY the translated text. Do NOT add explanations or notes.
#             5. If Name, surname, place of Birth already in Kazakh language, don't tranclate, 
#             Text to translate:

#             {text}

#             Fields to extract (JSON only):
#             1. Type (e.g., "P")
#             2. Code
#             3. Passport Number
#             4. Surname (original)
#             5. Given Names (original)
#             6. Nationality (translate to Russian if in Latin)
#             7. Date of Birth (dd.mm.yyyy)
#             8. Sex ("Ж" for F, "М" for M)
#             9. Authority (exact value)
#             10. ID Number
#             11. Place of Birth (translate to Russian if in Latin)
#             12. Date of Issue (dd.mm.yyyy)
#             13. Date of Expiry (dd.mm.yyyy)

#             If a field is missing, write "Not found".
#             """

#             response = openai.ChatCompletion.create(
#                 model="gpt-4",
#                 messages=[
#                     {"role": "system", "content": "You are a helpful assistant."},
#                     {"role": "user", "content": prompt}
#                 ]
#             )
#             result = re.sub(r'^.*?{', '{', result, flags=re.DOTALL)  # Удалить всё до первой {
#             result = re.sub(r'}.*?$', '}', result, flags=re.DOTALL)  # Удалить всё после последней }
#             result = response["choices"][0]["message"]["content"].strip()
#             try:
#                 # Удаляем потенциальные лишние строки
#                 start = result.find("{")
#                 end = result.rfind("}")
#                 if start != -1 and end != -1:
#                     result = result[start:end + 1]
#                     json.loads(result)  # Проверяем, что это корректный JSON
#                     return result
#             except json.JSONDecodeError:
#                 pass

#             print(f"Попытка {attempt + 1}: Некорректный JSON от OpenAI.")
#             time.sleep(2)  # Задержка перед повторной попыткой
#         except Exception as e:
#             print(f"Ошибка при обработке текста с помощью OpenAI (попытка {attempt + 1}): {e}")
#             time.sleep(2)  # Задержка перед повторной попыткой
            
#     print("OpenAI не смог вернуть корректный JSON после 3 попыток.")
#     return "{}"

# def is_already_translated(text):
#     """Проверяет, написан ли текст на нужном языке."""
#     cyrillic_chars = set("АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ")
#     return any(char in cyrillic_chars for char in text)

# def translate_field(field_name, field_value, target_language="ru"):
#     """Переводит значение поля через OpenAI."""
#     if re.search("[а-яА-ЯЁё]", field_value):
#         return field_value  #
#     try:
#         if is_already_translated(field_value):  # Новая функция для проверки
#             return field_value  # Возвращаем оригинальное значение

#         prompt = f"""
#         Translate the following into {target_language}:
#         {field_value}
#         """
#         response = openai.ChatCompletion.create(
#             model="gpt-4",
#             messages=[
#                 {"role": "system", "content": "You are a helpful assistant."},
#                 {"role": "user", "content": prompt}
#             ]
#         )
#         return response["choices"][0]["message"]["content"].strip()
#     except Exception as e:
#         print(f"Ошибка перевода для поля {field_name}: {e}")
#         return field_value

# def extract_and_fix_mrz_with_gpt(full_text):
#     """Извлекает и исправляет MRZ из полного текста с помощью OpenAI."""
#     try:
#         prompt = f"""
#         The following text is extracted from a passport image and contains the Machine Readable Zone (MRZ).
#         Your task is to:
#         1. Locate all lines that belong to the MRZ.
#         2. Merge them into exactly two lines, following the ICAO standard.
#         3. Ensure no data is lost during merging, even if the MRZ is split into more than two lines in the input text.
#         4. The MRZ should:
#             - Start with 'P<'.
#             - Have two lines.
#             - End the second line with a number.
#         5. Write only MRZ, without unneccerry information 

#         Text to process:
#         {full_text}

#         Provide only the corrected MRZ in two lines, ensuring all fields are present and properly formatted.
#         """
        
#         response = openai.ChatCompletion.create(
#             model="gpt-4",
#             messages=[
#                 {"role": "system", "content": "You are a helpful assistant that processes MRZ data from passport text."},
#                 {"role": "user", "content": prompt}
#             ]
#         )
        
#         # Получаем ответ GPT
#         fixed_mrz = response["choices"][0]["message"]["content"].strip()
#         return fixed_mrz
#     except Exception as e:
#         print(f"Ошибка при извлечении и исправлении MRZ с помощью GPT: {e}")
#         return "MRZ не найдено"

# def fill_word_template(template_path, output_path, data):
#     """Заполняет шаблон Word данными, включая текст внутри таблиц."""
#     doc = Document(template_path)

#     # Обработка текста в параграфах
#     for paragraph in doc.paragraphs:
#         for key, value in data.items():
#             if f"{{{{{key}}}}}" in paragraph.text:
#                 for run in paragraph.runs:
#                     if f"{{{{{key}}}}}" in run.text:
#                         run.text = run.text.replace(f"{{{{{key}}}}}", value.upper())

#     # Обработка текста внутри таблиц
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 # Обработка каждого параграфа внутри ячейки
#                 for paragraph in cell.paragraphs:
#                     for key, value in data.items():
#                         if f"{{{{{key}}}}}" in paragraph.text:
#                             for run in paragraph.runs:
#                                 if f"{{{{{key}}}}}" in run.text:
#                                     run.text = run.text.replace(f"{{{{{key}}}}}", value.upper())

#     # Сохранение документа
#     doc.save(output_path)
#     print(f"Данные сохранены в файл: {output_path}")

# def process_document(image_path, template_path, output_path):
#     """Полный процесс: OCR -> Перевод -> Заполнение Word шаблона."""
#     try:
#         # 1. Извлечение текста с помощью Google Vision API
#         print("Извлечение текста с изображения...")
#         recognized_text = extract_text_with_vision_api(image_path)
#         if not recognized_text:
#             print("Изображение не содержит текста.")
#             return

#         print("Распознанный текст:")
#         print(recognized_text)

#         # 2. Обработка текста с помощью OpenAI
#         print("Обработка текста с помощью OpenAI...")
#         processed_data = process_with_openai(recognized_text)
        
#         # Проверка и разбор JSON
#         try:
#             structured_data = json.loads(processed_data)
#         except json.JSONDecodeError as e:
#             print(f"Ошибка разбора JSON: {e}")
#             print("Ответ OpenAI:")
#             print(processed_data)
#             return

#         print("Извлечённые данные:")
#         print(json.dumps(structured_data, indent=4, ensure_ascii=False))

#         keys = ["Code", "Passport Number", "Surname", "Given Names", "Nationality", "Date of Birth",
#                 "Sex", "Place of Birth", "Date of Issue", "Date of Expiry", "Authority", "MRZ"]
#         for key in keys:
#             if key not in structured_data:
#                 structured_data[key] = "Not specified"
#         #3          
#         print("Извлечение и исправление MRZ с помощью GPT...")
#         fixed_mrz = extract_and_fix_mrz_with_gpt(recognized_text)
#         structured_data["MRZ"] = fixed_mrz
#         print(f"Исправленный MRZ: \n{fixed_mrz}")


#         print("Перевод данных:")
#         fields_to_translate = ["Surname", "Given Names", "Nationality", "Sex", "Place of Birth"]
#         for field in fields_to_translate:
#             if field in structured_data and structured_data[field] != "Not found":
#                 structured_data[field] = translate_field(field, structured_data[field], "ru")
#         print("Переведённые данные:")
#         print(json.dumps(structured_data, indent=4, ensure_ascii=False))

#         # 5. Заполнение шаблона Word
#         fill_word_template(template_path, output_path, structured_data)

#     except Exception as e:
#         print(f"Ошибка: {e}")

# if __name__ == "__main__":
#     image_path = "C:\\images\\page1_uzb.jpg"  # Путь к изображению
#     template_path = "C:\\images\\test_uzb.docx"  # Путь к шаблону Word
#     output_path = "C:\\images\\output_uzb.docx"
#     process_document(image_path, template_path, output_path)


# -----------------

import re
import json
import os
import time
from google.cloud import vision
from docx import Document
import openai
from dotenv import load_dotenv

# Загрузка API-ключей из .env файла
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

def extract_text_with_vision_api(image_path):
    """Распознаёт текст с изображения с помощью Google Vision API."""
    print("\n🔍 Извлекаем текст с изображения...")
    client = vision.ImageAnnotatorClient()
    
    with open(image_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if not texts:
        print("❌ Текст не найден на изображении.")
        return ""

    extracted_text = texts[0].description.strip()
    print(f"✅ Извлечённый текст:\n{extracted_text[:500]}...")
    return extracted_text

def process_with_openai(text):
    """Обрабатывает текст с помощью OpenAI (извлекает паспортные данные)."""
    print("\n📡 Отправляем данные в OpenAI для обработки...")
    prompt = f"""
    You are processing a Kazakh passport. Return only JSON without any extra explanations.
    
    - Names (Surname, Given Names) **MUST BE TRANSLATED** to Russian.
    - Geographic names (Place of Birth, Authority) **MUST BE TRANSLATED**.
    - Nationality must be translated if written in Latin script.
    - DO NOT ADD EXTRA COMMENTS OR EXPLANATIONS.

    Extract the following data as JSON:
    - Type
    - Code
    - Passport Number
    - Surname (translate to Russian)
    - Given Names (translate to Russian)
    - Nationality (translate if in Latin)
    - Date of Birth (dd.mm.yyyy)
    - Sex ("Ж" for F, "М" for M)
    - Authority (translate)
    - Place of Birth (translate)
    - Date of Issue (dd.mm.yyyy)
    - Date of Expiry (dd.mm.yyyy)

    Text to process:
    {text}
    """

    for _ in range(3):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[{"role": "system", "content": "You are a professional AI that processes passports."},
                          {"role": "user", "content": prompt}],
                response_format={"type": "json_object"}
            )
            structured_data = json.loads(response["choices"][0]["message"]["content"])
            print(f"✅ OpenAI вернул данные:\n{json.dumps(structured_data, indent=4, ensure_ascii=False)}")
            return structured_data
        except Exception as e:
            print(f"⚠️ Ошибка OpenAI: {e}")
            time.sleep(2)
    return {}

def fill_word_template(template_path, output_path, data):
    """Заполняет шаблон Word данными и удаляет placeholders."""
    print("\n📝 Заполняем шаблон Word...")
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{{{key}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value.upper())

    # Убираем placeholders, если данные не найдены
    for paragraph in doc.paragraphs:
        paragraph.text = re.sub(r'\{\{.*?\}\}', '', paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = re.sub(r'\{\{.*?\}\}', '', paragraph.text)

    doc.save(output_path)
    print(f"✅ Документ сохранён: {output_path}")

def process_document(image_path, template_path, output_path):
    """Основной процесс: OCR -> обработка -> перевод -> заполнение шаблона."""
    try:
        recognized_text = extract_text_with_vision_api(image_path)
        if not recognized_text:
            print("❌ Изображение не содержит текста.")
            return
        
        structured_data = process_with_openai(recognized_text)

        fill_word_template(template_path, output_path, structured_data)

        print("\n🎉 Процесс завершён успешно!")
    except Exception as e:
        print(f"❌ Ошибка: {e}")

if __name__ == "__main__":
    process_document("input.jpg", "template.docx", "output.docx")
